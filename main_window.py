import os
import sys
import cv2
import time
import numpy as np
import json
import threading
from collections import defaultdict, deque
from datetime import datetime

import torch
from ultralytics import YOLO, YOLOWorld

from PyQt5.QtWidgets import (
    QWidget, QLabel, QPushButton, QVBoxLayout, QHBoxLayout, QLineEdit,
    QFileDialog, QComboBox, QMessageBox, QTableWidget,
    QTableWidgetItem, QHeaderView, QSlider, QFrame, QGroupBox,
    QSpacerItem, QSizePolicy, QCheckBox, QSplitter, QListWidget,
    QListWidgetItem, QMenu, QAction, QInputDialog, QProgressDialog,
    QShortcut
)
from PyQt5.QtGui import (
    QPixmap, QImage, QFont, QColor, QKeySequence
)
from PyQt5.QtCore import (
    Qt, QTimer, QMutex, QMutexLocker, QRectF, QPointF,
    QThread, QRect, QSize, QPoint
)

from config import db_connect
from theme import STYLE_SHEET
from traffic_monitor import TrafficMonitor
from widgets import EditableImageLabel
from qwen_analyzer import QwenAnalyzer
from history import HistoryWindow

_BASE_DIR = os.path.dirname(os.path.abspath(__file__))


# ---------- 主检测窗口 ----------
class YOLOv10App(QWidget):
    def __init__(self, username):
        super().__init__()
        self.username = username
        self.track_history = defaultdict(lambda: [])
        self.class_stats = defaultdict(int)
        self.use_ovd = False
        self.ovd_labels = []
        self.detection_mode = "standard"
        self.intrusion_zones = []
        self.intrusion_alerts = deque(maxlen=10)
        self.enable_privacy = False
        self.active_learning_enabled = True
        self.show_alert_banner = True  # 画面预警图标开关
        self.correction_queue = []
        self.frame_count = 0  # 帧计数器，用于交错优化

        # 交通监控
        self.traffic_monitor = TrafficMonitor()
        self.last_alerts = []
        self.alert_history = []
        self.qwen_analyzer = QwenAnalyzer(api_key="sk-3266b6cdde9e492fbb10cdda6b3e9e7f")

        self.setWindowTitle(f"检测系统 · {self.username} [前沿版]")
        self.setMinimumSize(1400, 860)
        self.resize(1600, 920)
        self.setObjectName("centralWidget")
        self.setStyleSheet(STYLE_SHEET)

        # 模型加载
        try:
            self.model = YOLO(os.path.join(_BASE_DIR, 'yolov10s.pt'))
            self.ovd_model = None
            try:
                self.ovd_model = YOLOWorld(os.path.join(_BASE_DIR, 'yolov8s-world.pt'))  # YOLOv10无World变体，OVD使用World模型
                print("[模型] 开放词汇模型加载成功")
            except:
                print("[模型] 开放词汇模型不可用")
        except Exception as e:
            QMessageBox.critical(self, "模型错误", f"无法加载 YOLOv10 模型：{str(e)}")
            sys.exit(1)

        self.cap = None
        self.timer = QTimer()
        self.mutex = QMutex()
        self.confidence_threshold = 0.5

        # ===== 主布局 =====
        main_layout = QHBoxLayout()
        main_layout.setContentsMargins(16, 16, 16, 16)
        main_layout.setSpacing(16)

        # 左侧：图像区域
        left_widget = QWidget()
        left_layout = QVBoxLayout(left_widget)
        left_layout.setContentsMargins(0, 0, 0, 0)
        left_layout.setSpacing(12)

        header_frame = QFrame()
        header_frame.setObjectName("cardDark")
        header_frame.setStyleSheet("""
            QFrame#cardDark {
                background-color: #0f172a;
                border-radius: 16px;
            }
        """)
        header_layout = QHBoxLayout(header_frame)
        header_layout.setContentsMargins(20, 12, 20, 12)
        brand_label = QLabel("YOLOv10s · 目标检测与追踪 · 交通预警")
        brand_label.setStyleSheet("color: #f1f5f9; font-size: 14px; font-weight: 500; letter-spacing: -0.2px;")
        header_layout.addWidget(brand_label)
        header_layout.addStretch()
        version_label = QLabel("v3.0")
        version_label.setStyleSheet("color: #475569; font-size: 12px; margin-right: 8px;")
        header_layout.addWidget(version_label)
        user_info = QLabel(self.username)
        user_info.setStyleSheet("color: #94a3b8; font-size: 13px; font-weight: 500;")
        header_layout.addWidget(user_info)

        self.image_frame = QFrame()
        self.image_frame.setObjectName("card")
        self.image_frame.setStyleSheet("""
            QFrame#card {
                background-color: #ffffff;
                border: 1px solid #f1f5f9;
                border-radius: 16px;
            }
        """)
        img_layout = QVBoxLayout(self.image_frame)
        img_layout.setContentsMargins(10, 10, 10, 10)

        self.image_label = EditableImageLabel()
        self.image_label.setAlignment(Qt.AlignCenter)
        self.image_label.setMinimumHeight(400)
        self.image_label.setStyleSheet("""
            background-color: #fafbfc;
            border-radius: 14px;
            color: #94a3b8;
            font-size: 15px;
        """)
        # 连接自定义信号
        self.image_label.rect_created.connect(self.add_intrusion_zone_from_rect)
        self.image_label.line_created.connect(self.add_direction_line)
        self.image_label.setContextMenuPolicy(Qt.CustomContextMenu)
        self.image_label.customContextMenuRequested.connect(self.image_context_menu)
        img_layout.addWidget(self.image_label)

        self.status_label = QLabel("系统就绪")
        self.status_label.setAlignment(Qt.AlignCenter)
        self.status_label.setStyleSheet("color: #64748b; font-size: 13px; padding: 6px 16px; background: #ffffff; border-radius: 10px;")

        left_layout.addWidget(header_frame)
        left_layout.addWidget(self.image_frame, 1)
        left_layout.addWidget(self.status_label)

        # 右侧面板
        right_widget = QWidget()
        right_layout = QVBoxLayout(right_widget)
        right_layout.setSpacing(12)
        right_layout.setContentsMargins(0, 0, 0, 0)

        # 检测设置
        settings_group = QGroupBox("检测设置")
        settings_layout = QVBoxLayout()
        settings_layout.setSpacing(12)

        mode_layout = QHBoxLayout()
        mode_layout.addWidget(QLabel("模式:"))
        self.mode_combo = QComboBox()
        self.mode_combo.addItems(["标准检测", "开放词汇"])
        self.mode_combo.currentIndexChanged.connect(self.switch_mode)
        mode_layout.addWidget(self.mode_combo)
        settings_layout.addLayout(mode_layout)

        self.ovd_input = QLineEdit()
        self.ovd_input.setPlaceholderText("输入自然语言标签，如：红色安全帽，叉车")
        self.ovd_input.setVisible(False)
        self.ovd_input.returnPressed.connect(self.apply_ovd_labels)
        settings_layout.addWidget(self.ovd_input)

        self.target_selector = QComboBox()
        target_items = ["全部目标"] + [str(cname) for cname in self.model.names.values()]
        self.target_selector.addItems(target_items)
        settings_layout.addWidget(self.target_selector)

        conf_layout = QVBoxLayout()
        conf_layout.setSpacing(6)
        self.conf_label = QLabel("置信度阈值: 0.50")
        self.conf_label.setAlignment(Qt.AlignCenter)
        self.conf_label.setStyleSheet("font-size: 13px; color: #64748b;")
        self.conf_slider = QSlider(Qt.Horizontal)
        self.conf_slider.setRange(0, 100)
        self.conf_slider.setValue(50)
        self.conf_slider.valueChanged.connect(self.update_confidence)
        conf_layout.addWidget(self.conf_label)
        conf_layout.addWidget(self.conf_slider)

        self.privacy_cb = QCheckBox("隐私保护（人脸模糊）")
        self.privacy_cb.toggled.connect(lambda v: setattr(self, 'enable_privacy', v))
        self.active_learn_cb = QCheckBox("启用主动学习")
        self.active_learn_cb.setChecked(True)
        self.active_learn_cb.toggled.connect(lambda v: setattr(self, 'active_learning_enabled', v))
        settings_layout.addLayout(conf_layout)
        settings_layout.addWidget(self.privacy_cb)
        settings_layout.addWidget(self.active_learn_cb)
        settings_group.setLayout(settings_layout)

        # 交通预警功能区（新增）
        self.alert_group = QGroupBox("🚦 自动预警（始终开启）")
        self.alert_summary = QLabel("暂无预警")
        self.alert_summary.setStyleSheet("color: #64748b; font-size: 12px; padding: 2px 0;")
        alert_layout = QVBoxLayout()
        self.alert_list = QListWidget()
        self.alert_list.setAlternatingRowColors(True)
        self.alert_list.setStyleSheet("""
            QListWidget::item { padding: 6px; border-bottom: 1px solid #f1f5f9; }
        """)
        self.alert_banner_cb = QCheckBox("📺 画面直接预警")
        self.alert_banner_cb.setChecked(True)
        self.alert_banner_cb.setToolTip("开启后预警信息将直接绘制在视频画面上")
        self.alert_banner_cb.toggled.connect(self.toggle_alert_banner)

        self.clear_alerts_btn = QPushButton("清除预警")
        self.clear_alerts_btn.clicked.connect(self.clear_alerts)

        # ---- 手动设置：类型 & 等级过滤 ----
        filter_type_layout = QHBoxLayout()
        filter_type_layout.setSpacing(6)
        filter_type_layout.addWidget(QLabel("类型:"))
        self.filter_type_cbs = {}
        for at in ["逆行", "行人闯入", "车辆碰撞", "区域入侵"]:
            cb = QCheckBox(at)
            cb.setChecked(True)
            cb.toggled.connect(self._apply_alert_filter)
            self.filter_type_cbs[at] = cb
            filter_type_layout.addWidget(cb)
        filter_type_layout.addStretch()

        filter_level_layout = QHBoxLayout()
        filter_level_layout.setSpacing(6)
        filter_level_layout.addWidget(QLabel("等级:"))
        self.filter_level_cbs = {}
        for lv, name in [(3, "危险"), (2, "警告"), (1, "注意")]:
            cb = QCheckBox(name)
            cb.setChecked(True)
            cb.toggled.connect(self._apply_alert_filter)
            self.filter_level_cbs[lv] = cb
            filter_level_layout.addWidget(cb)
        filter_level_layout.addStretch()

        # ---- 手动触发按钮 ----
        self.manual_alert_btn = QPushButton("手动触发预警")
        self.manual_alert_btn.setObjectName("manualAlertBtn")
        self.manual_alert_btn.setMinimumHeight(38)
        self.manual_alert_btn.setCursor(Qt.PointingHandCursor)
        self.manual_alert_btn.clicked.connect(self.manual_trigger_alert)

        alert_layout.addWidget(self.alert_summary)
        alert_layout.addLayout(filter_type_layout)
        alert_layout.addLayout(filter_level_layout)
        alert_layout.addWidget(self.alert_list)
        alert_layout.addWidget(self.alert_banner_cb)
        alert_layout.addWidget(self.clear_alerts_btn)
        alert_layout.addWidget(self.manual_alert_btn)
        self.alert_group.setLayout(alert_layout)

        # 统计卡片
        stats_group = QGroupBox("实时统计")
        stats_layout = QVBoxLayout()
        self.stats_table = QTableWidget()
        self.stats_table.setColumnCount(2)
        self.stats_table.setHorizontalHeaderLabels(["目标", "数量"])
        self.stats_table.horizontalHeader().setSectionResizeMode(QHeaderView.Stretch)
        self.stats_table.setAlternatingRowColors(True)
        self.stats_table.verticalHeader().setVisible(False)
        stats_layout.addWidget(self.stats_table)
        stats_group.setLayout(stats_layout)

        # 操作卡片
        action_group = QGroupBox("操作")
        action_layout = QVBoxLayout()
        action_layout.setSpacing(8)

        self.img_button = QPushButton("检测图片")
        self.vid_button = QPushButton("检测视频")
        self.cam_button = QPushButton("摄像头检测")
        self.stop_button = QPushButton("停止检测")
        self.stop_button.setObjectName("stopBtn")
        self.stop_button.setEnabled(False)
        self.history_button = QPushButton("检测记录")
        self.report_button = QPushButton("生成报告")

        for btn in [self.img_button, self.vid_button, self.cam_button,
                     self.stop_button, self.history_button, self.report_button]:
            btn.setMinimumHeight(38)
            btn.setCursor(Qt.PointingHandCursor)

        self.img_button.clicked.connect(self.select_image)
        self.vid_button.clicked.connect(self.select_video)
        self.cam_button.clicked.connect(self.detect_camera)
        self.stop_button.clicked.connect(self.stop_detection)
        self.history_button.clicked.connect(self.show_history)
        self.report_button.clicked.connect(self.generate_report)

        action_layout.addWidget(self.img_button)
        action_layout.addWidget(self.vid_button)
        action_layout.addWidget(self.cam_button)
        action_layout.addWidget(self.stop_button)
        action_layout.addWidget(self.history_button)
        action_layout.addWidget(self.report_button)
        action_group.setLayout(action_layout)

        # 右侧面板组装
        right_layout.addWidget(settings_group)
        right_layout.addWidget(self.alert_group)   # 预警功能区
        right_layout.addWidget(stats_group)
        right_layout.addWidget(action_group)
        right_layout.addStretch()

        main_layout.addWidget(left_widget, 7)
        main_layout.addWidget(right_widget, 3)
        self.setLayout(main_layout)


        # 快捷键
        QShortcut(QKeySequence("Ctrl+O"), self, self.select_image)
        QShortcut(QKeySequence("Ctrl+V"), self, self.select_video)
        QShortcut(QKeySequence("Ctrl+C"), self, self.detect_camera)
        QShortcut(QKeySequence("Ctrl+S"), self, self.stop_detection)
        QShortcut(QKeySequence("Ctrl+H"), self, self.show_history)

    # ---------- 模式切换 ----------
    def switch_mode(self, index):
        if index == 0:
            self.detection_mode = "standard"
            self.target_selector.setVisible(True)
            self.ovd_input.setVisible(False)
        else:
            self.detection_mode = "ovd"
            self.target_selector.setVisible(False)
            self.ovd_input.setVisible(True)

    def apply_ovd_labels(self):
        text = self.ovd_input.text().strip()
        if text:
            self.ovd_labels = [t.strip() for t in text.split(",") if t.strip()]
            if self.ovd_model and self.ovd_labels:
                self.ovd_model.set_classes(self.ovd_labels)
            self.status_label.setText(f"开放词汇标签已设置: {', '.join(self.ovd_labels)}")

    def update_confidence(self):
        self.confidence_threshold = self.conf_slider.value() / 100
        self.conf_label.setText(f"置信度阈值: {self.confidence_threshold:.2f}")

    # ---------- 交通规则配置 ----------
    def add_intrusion_zone_from_rect(self, rect):
        """来自图像标签的矩形（入侵区域），同时加入 TrafficMonitor 的禁行区"""
        zone = (rect.x(), rect.y(), rect.right(), rect.bottom())
        self.intrusion_zones.append(zone)
        self.traffic_monitor.set_restricted_zones(self.intrusion_zones)
        self.status_label.setText(f"入侵/禁行区已添加：{zone}")

    def add_direction_line(self, start, end):
        """来自图像标签的方向线，加入 TrafficMonitor 的允许方向"""
        self.traffic_monitor.direction_lines.append((start, end))
        self.status_label.setText(f"方向线已添加：({start.x():.0f},{start.y():.0f}) -> ({end.x():.0f},{end.y():.0f})")

    def image_context_menu(self, pos):
        menu = QMenu()
        menu.addAction("清除入侵区域/禁行区").triggered.connect(lambda: self.clear_intrusion_zones())
        menu.addAction("清除所有方向线").triggered.connect(lambda: self.clear_direction_lines())
        menu.exec_(self.image_label.mapToGlobal(pos))

    def clear_intrusion_zones(self):
        self.intrusion_zones.clear()
        self.image_label.rois.clear()
        self.image_label.update()
        self.traffic_monitor.set_restricted_zones([])
        self.status_label.setText("入侵/禁行区已清除")

    def clear_direction_lines(self):
        self.traffic_monitor.direction_lines.clear()
        self.image_label.lines.clear()
        self.image_label.update()
        self.status_label.setText("方向线已清除")

    def toggle_alert_banner(self, enabled):
        """手动开关画面预警图标"""
        self.show_alert_banner = enabled
        if enabled:
            self.status_label.setText("画面预警图标：已开启")
        else:
            self.status_label.setText("画面预警图标：已关闭，仅侧边栏提醒")

    def clear_alerts(self):
        self.alert_list.clear()
        self.last_alerts.clear()
        self.alert_history.clear()
        self.traffic_monitor.alert_cooldown.clear()
        self.alert_summary.setText("暂无预警")
        self.alert_group.setTitle("🚦 自动预警（始终开启）")

    def manual_trigger_alert(self):
        """手动触发预警：直接往预警面板灌入演示数据"""
        import time as _time
        now = _time.time()

        demo_alerts = [
            (12, "car↔truck",   "车辆碰撞", 3, now - 5),
            (8,  "motorcycle",   "逆行",     2, now - 12),
            (3,  "person",       "行人闯入", 2, now - 25),
            (15, "person",       "区域入侵", 2, now - 38),
            (6,  "bicycle",      "逆行",     1, now - 50),
            (9,  "person",       "行人闯入", 1, now - 65),
        ]

        for (obj_id, cls_label, alert_type, level, timestamp) in demo_alerts:
            time_str = datetime.fromtimestamp(timestamp).strftime("%H:%M:%S")
            self.alert_history.append({
                "obj_id": obj_id, "label": cls_label,
                "type": alert_type, "level": level,
                "level_name": {3: "危险", 2: "警告", 1: "注意"}[level],
                "time": time_str, "timestamp": timestamp
            })

        # 同时更新 traffic_monitor cooldown，防止检测线程立即覆盖
        self._apply_alert_filter()
        self.status_label.setText("预警已手动触发，共 {} 条".format(len(self.alert_history)))

    def _apply_alert_filter(self):
        """根据手动设置的类型和等级过滤并刷新预警列表"""
        selected_types = {t for t, cb in self.filter_type_cbs.items() if cb.isChecked()}
        selected_levels = {lv for lv, cb in self.filter_level_cbs.items() if cb.isChecked()}

        filtered = [
            a for a in self.alert_history
            if a.get("type") in selected_types and a.get("level") in selected_levels
        ]

        self.alert_list.clear()
        level_color = {1: "#3b82f6", 2: "#f59e0b", 3: "#ef4444"}
        level_text = {1: "注意", 2: "警告", 3: "危险"}

        if filtered:
            for a in reversed(filtered[-30:]):
                text = f"[{a.get('level_name','?')}] {a.get('type','')} — {a.get('label','')}-{a.get('obj_id','')}    {a.get('time','')}"
                item = QListWidgetItem(text)
                item.setForeground(QColor(level_color.get(a.get("level", 1), "#000")))
                self.alert_list.addItem(item)
        else:
            placeholder = QListWidgetItem("（无匹配预警）")
            placeholder.setForeground(QColor("#94a3b8"))
            self.alert_list.addItem(placeholder)

        # 更新统计
        total = len(self.alert_history)
        ft = len(filtered)
        l3 = sum(1 for a in filtered if a.get("level") == 3)
        l2 = sum(1 for a in filtered if a.get("level") == 2)
        l1 = sum(1 for a in filtered if a.get("level") == 1)

        if total > 0:
            self.alert_summary.setText(f"显示 {ft}/{total} 条 | 危险 {l3} · 警告 {l2} · 注意 {l1}")
            if l3 > 0:
                self.alert_group.setTitle("🚦 预警 ⚠ 危险")
            elif l2 > 0:
                self.alert_group.setTitle("🚦 预警 ⚡ 警告")
            else:
                self.alert_group.setTitle("🚦 预警 ℹ 注意")
        else:
            self.alert_summary.setText("暂无预警")
            self.alert_group.setTitle("🚦 自动预警")

    # ---------- 类别名称映射 ----------
    # COCO 标准 80 类名称，作为兜底映射
    COCO_NAMES = {
        0: 'person', 1: 'bicycle', 2: 'car', 3: 'motorcycle', 4: 'airplane', 5: 'bus',
        6: 'train', 7: 'truck', 8: 'boat', 9: 'traffic light', 10: 'fire hydrant',
        11: 'stop sign', 12: 'parking meter', 13: 'bench', 14: 'bird', 15: 'cat',
        16: 'dog', 17: 'horse', 18: 'sheep', 19: 'cow', 20: 'elephant', 21: 'bear',
        22: 'zebra', 23: 'giraffe', 24: 'backpack', 25: 'umbrella', 26: 'handbag',
        27: 'tie', 28: 'suitcase', 29: 'frisbee', 30: 'skis', 31: 'snowboard',
        32: 'sports ball', 33: 'kite', 34: 'baseball bat', 35: 'baseball glove',
        36: 'skateboard', 37: 'surfboard', 38: 'tennis racket', 39: 'bottle',
        40: 'wine glass', 41: 'cup', 42: 'fork', 43: 'knife', 44: 'spoon', 45: 'bowl',
        46: 'banana', 47: 'apple', 48: 'sandwich', 49: 'orange', 50: 'broccoli',
        51: 'carrot', 52: 'hot dog', 53: 'pizza', 54: 'donut', 55: 'cake',
        56: 'chair', 57: 'couch', 58: 'potted plant', 59: 'bed', 60: 'dining table',
        61: 'toilet', 62: 'tv', 63: 'laptop', 64: 'mouse', 65: 'remote',
        66: 'keyboard', 67: 'cell phone', 68: 'microwave', 69: 'oven', 70: 'toaster',
        71: 'sink', 72: 'refrigerator', 73: 'book', 74: 'clock', 75: 'vase',
        76: 'scissors', 77: 'teddy bear', 78: 'hair drier', 79: 'toothbrush',
    }

    def _class_name(self, cls_id):
        """将类别编号(0,1,2...)映射为对应的识别类型英语单词(person, car, truck...)"""
        cid = int(cls_id)
        # 优先从模型获取
        try:
            raw = self.model.names[cid]
        except (KeyError, IndexError, TypeError):
            raw = None
        # 如果模型返回的是数字或无效值，用 COCO 兜底
        name = str(raw) if raw is not None and str(raw).strip() else ""
        if not name or name.isdigit():
            name = self.COCO_NAMES.get(cid, "")
        return name if name else f"class_{cid}"

    def _tiled_detect(self, frame, model, tile_rows=2, tile_cols=2, overlap=0.20, is_image=False):
        """分块检测增强：2x2重叠切片 + 水平翻转，每4帧触发一次"""
        h, w = frame.shape[:2]
        tile_h = int(h / tile_rows)
        tile_w = int(w / tile_cols)
        overlap_h = int(tile_h * overlap)
        overlap_w = int(tile_w * overlap)

        all_boxes = []
        vehicle_ids = {2, 3, 5, 7, 8, 4, 6}
        det_conf = max(0.15, self.confidence_threshold * 0.5)

        for row in range(tile_rows):
            for col in range(tile_cols):
                y1 = max(0, row * tile_h - overlap_h)
                y2 = min(h, (row + 1) * tile_h + overlap_h)
                x1 = max(0, col * tile_w - overlap_w)
                x2 = min(w, (col + 1) * tile_w + overlap_w)

                tile = frame[y1:y2, x1:x2]
                if tile.size == 0 or tile.shape[0] < 30 or tile.shape[1] < 30:
                    continue

                # ---- 原方向检测 ----
                self._detect_tile(tile, x1, y1, w, h, all_boxes, model, det_conf, vehicle_ids, flipped=False, is_image=is_image)

                # ---- 水平翻转检测（车头朝左/右都能识别） ----
                tile_flipped = cv2.flip(tile, 1)
                self._detect_tile(tile_flipped, x1, y1, w, h, all_boxes, model, det_conf, vehicle_ids, flipped=True, is_image=is_image)

        if not all_boxes:
            return None

        keep = self._nms(all_boxes, iou_thresh=0.5)
        return [all_boxes[i] for i in keep]

    def _detect_tile(self, tile, x1, y1, frame_w, frame_h, all_boxes, model, det_conf, vehicle_ids, flipped, is_image=False):
        """对单个tile执行检测，flipped=True时自动镜像坐标"""
        det_func = model.predict if is_image else model.track
        tile_results = det_func(tile, persist=False, conf=det_conf,
                                imgsz=960, iou=0.5, verbose=False)
        if tile_results[0].boxes is None or tile_results[0].boxes.xyxy is None:
            return

        tw, th = tile.shape[1], tile.shape[0]
        for box in tile_results[0].boxes:
            cls_id = int(box.cls[0])
            if cls_id not in vehicle_ids:
                continue
            conf = float(box.conf[0])
            xyxy = box.xyxy[0].cpu().tolist()

            if flipped:
                # 水平翻转坐标映射：x' = tile_w - x
                x_left  = tw - xyxy[2]
                x_right = tw - xyxy[0]
                xyxy[0] = x_left
                xyxy[2] = x_right

            # 映射回原图
            xyxy[0] += x1; xyxy[1] += y1
            xyxy[2] += x1; xyxy[3] += y1

            # 半截入镜提权
            edge_touch = (xyxy[0] <= 2 or xyxy[1] <= 2 or
                          xyxy[2] >= frame_w - 2 or xyxy[3] >= frame_h - 2)
            if edge_touch:
                conf = min(1.0, conf * 1.4)

            all_boxes.append({'cls': cls_id, 'conf': conf, 'xyxy': xyxy})

    @staticmethod
    def _nms(boxes, iou_thresh=0.45):
        """NMS去重，返回保留的索引列表"""
        if not boxes:
            return []
        indices = sorted(range(len(boxes)), key=lambda i: boxes[i]['conf'], reverse=True)
        keep = []
        while indices:
            current = indices.pop(0)
            keep.append(current)
            cb = boxes[current]['xyxy']
            filtered = []
            for i in indices:
                ob = boxes[i]['xyxy']
                xA = max(cb[0], ob[0]); yA = max(cb[1], ob[1])
                xB = min(cb[2], ob[2]); yB = min(cb[3], ob[3])
                inter = max(0, xB - xA) * max(0, yB - yA)
                area_a = (cb[2] - cb[0]) * (cb[3] - cb[1])
                area_b = (ob[2] - ob[0]) * (ob[3] - ob[1])
                iou = inter / (area_a + area_b - inter + 1e-6)
                if iou < iou_thresh:
                    filtered.append(i)
            indices = filtered
        return keep

    # ---------- 核心检测引擎 ----------
    def detect_objects(self, frame, is_image=False):
        self.class_stats.clear()
        if self.enable_privacy:
            frame = self.apply_privacy_mask(frame)

        self.frame_count += 1
        model = self.ovd_model if (self.detection_mode == "ovd" and self.ovd_model and self.ovd_labels) else self.model

        frame_h, frame_w = frame.shape[:2]
        vehicle_ids = {2, 3, 5, 7, 8, 4, 6}

        # 主检测：图片模式用更高分辨率，视频模式平衡速度与精度
        imgsz = 1280 if is_image else 960
        det_conf = max(0.15, self.confidence_threshold * 0.6) if is_image else self.confidence_threshold

        if is_image:
            results = model.predict(frame, conf=det_conf, imgsz=imgsz, iou=0.45, verbose=False)
        else:
            results = model.track(frame, persist=True, conf=det_conf,
                                  imgsz=imgsz, iou=0.45, verbose=False)

        # 收集已有车辆框
        existing_boxes = []  # [(x1,y1,x2,y2), ...]
        if results[0].boxes is not None:
            for box in results[0].boxes:
                if int(box.cls[0]) in vehicle_ids:
                    existing_boxes.append(box.xyxy[0].cpu().tolist())

        # ---- 水平翻转增强：图片模式始终执行，视频模式每 2 帧 ----
        run_flip = is_image or (self.frame_count % 2 == 0)
        if run_flip:
            flipped_frame = cv2.flip(frame, 1)
            flip_conf = det_conf * 0.7
            flip_results = det_func(flipped_frame, persist=False,
                                    conf=flip_conf,
                                    imgsz=imgsz, iou=0.45, verbose=False)
            if flip_results[0].boxes is not None and flip_results[0].boxes.xyxy is not None:
                for box in flip_results[0].boxes:
                    cls_id = int(box.cls[0])
                    if cls_id not in vehicle_ids:
                        continue
                    xyxy = box.xyxy[0].cpu().tolist()
                    x_left  = frame_w - xyxy[2]
                    x_right = frame_w - xyxy[0]
                    xyxy[0] = x_left
                    xyxy[2] = x_right
                    existing_boxes.append(xyxy)

        # ---- 分块检测增强：图片模式始终执行，视频模式每 4 帧 ----
        extra_vehicle_boxes = []
        if self.detection_mode == "standard" and (is_image or self.frame_count % 4 == 0):
            tiled_boxes = self._tiled_detect(frame, model, is_image=is_image)
            if tiled_boxes:
                for tb in tiled_boxes:
                    is_dup = False
                    for eb in existing_boxes:
                        xA = max(tb['xyxy'][0], eb[0]); yA = max(tb['xyxy'][1], eb[1])
                        xB = min(tb['xyxy'][2], eb[2]); yB = min(tb['xyxy'][3], eb[3])
                        inter = max(0, xB - xA) * max(0, yB - yA)
                        area_t = (tb['xyxy'][2]-tb['xyxy'][0]) * (tb['xyxy'][3]-tb['xyxy'][1])
                        area_e = (eb[2]-eb[0]) * (eb[3]-eb[1])
                        iou = inter / (area_t + area_e - inter + 1e-6)
                        if iou > 0.45:
                            is_dup = True
                            break
                    if not is_dup:
                        extra_vehicle_boxes.append(tb)

        # 交通规则同步
        self.traffic_monitor.set_restricted_zones(self.intrusion_zones)

        new_alerts = []
        all_vehicle_boxes = []

        # ---- 处理主检测结果 ----
        if results[0].boxes is not None:
            boxes_data = results[0].boxes
            # predict() 无 tracking ID，用序号代替；track() 有 .id
            if hasattr(boxes_data, 'id') and boxes_data.id is not None:
                box_ids = boxes_data.id.int().cpu().tolist()
            else:
                box_ids = list(range(len(boxes_data)))
            for box, box_id in zip(boxes_data, box_ids):
                cls_id = int(box.cls[0])
                label = self._class_name(cls_id) if self.detection_mode == "standard" else (
                    self.ovd_labels[cls_id] if cls_id < len(self.ovd_labels) else "unknown")

                selected_target = self.target_selector.currentText()
                if self.detection_mode == "standard" and selected_target != "全部目标" and label != selected_target:
                    continue

                x1, y1, x2, y2 = map(int, box.xyxy[0])
                self.class_stats[label] += 1
                all_vehicle_boxes.append((int(box_id), label, (x1, y1, x2, y2)))

                in_zone, zone_alert = self.check_intrusion(x1, y1, x2, y2, box_id, label)
                color = (0, 0, 255) if in_zone else (59, 130, 246)
                if zone_alert:
                    new_alerts.append(zone_alert)

                cv2.rectangle(frame, (x1, y1), (x2, y2), color, 3)
                cv2.putText(frame, f'{label} #{box_id}', (x1, y1 - 8),
                            cv2.FONT_HERSHEY_SIMPLEX, 0.7, color, 2)

                x_center = (x1 + x2) // 2
                y_center = (y1 + y2) // 2
                self.track_history[box_id].append((x_center, y_center))
                if len(self.track_history[box_id]) > 30:
                    self.track_history[box_id].pop(0)
                if len(self.track_history[box_id]) > 1:
                    points = np.array(self.track_history[box_id], dtype=np.int32)
                    cv2.polylines(frame, [points], isClosed=False, color=(250, 204, 21), thickness=3)

                traffic_alerts = self.traffic_monitor.update(box_id, label, (x_center, y_center))
                new_alerts.extend(traffic_alerts)

        # ---- 处理分块检测补充的车辆 ----
        fake_id = 90000  # 用高位ID避免与追踪ID冲突
        for tb in extra_vehicle_boxes:
            cls_id = tb['cls']
            label = self._class_name(cls_id)
            x1, y1, x2, y2 = map(int, tb['xyxy'])
            bid = fake_id
            fake_id += 1

            # 检查目标过滤器
            selected_target = self.target_selector.currentText()
            if selected_target != "全部目标" and label != selected_target:
                continue

            self.class_stats[label] += 1
            all_vehicle_boxes.append((bid, label, (x1, y1, x2, y2)))

            # 分块补充的用橙色框区分
            cv2.rectangle(frame, (x1, y1), (x2, y2), (0, 165, 255), 2)
            cv2.putText(frame, f'{label} #{bid}', (x1, y1 - 8),
                        cv2.FONT_HERSHEY_SIMPLEX, 0.55, (0, 165, 255), 2)

        # ---- 碰撞检测 ----
        collision_alerts = self.traffic_monitor.check_collisions(all_vehicle_boxes)
        new_alerts.extend(collision_alerts)
        for alert in collision_alerts:
            parts = alert[1].split('↔')
            for part in parts:
                for box_info in all_vehicle_boxes:
                    if str(box_info[0]) == part:
                        x1, y1, x2, y2 = box_info[2]
                        cv2.rectangle(frame, (x1-3, y1-3), (x2+3, y2+3), (0, 0, 255), 4)
                        cv2.putText(frame, 'COLLISION', (x1, y1 - 30),
                                    cv2.FONT_HERSHEY_SIMPLEX, 0.9, (0, 0, 255), 3)
                        break

        # ---- 更新UI ----
        self.update_alert_ui(new_alerts)
        if self.show_alert_banner:
            self.draw_alert_banner(frame, new_alerts)
        self.update_stats_table()
        return frame

    def draw_alert_banner(self, frame, new_alerts):
        """在视频画面上直接绘制醒目的预警信息（cv2只支持ASCII，全部用英文）"""
        if not new_alerts:
            return

        h, w = frame.shape[:2]
        level_color_map = {
            1: ((59, 130, 246),   "#3b82f6"),
            2: ((245, 158, 11),   "#f59e0b"),
            3: ((239, 68, 68),    "#ef4444"),
        }
        level_text_map = {1: "Caution", 2: "WARNING", 3: "DANGER"}
        # cv2不能画中文，将预警类型映射为英文
        type_en = {"逆行": "Wrong Way", "行人闯入": "Pedestrian", "车辆碰撞": "Collision", "区域入侵": "Zone Intrusion"}

        # 按等级从高到低排序
        sorted_alerts = sorted(new_alerts, key=lambda a: a[3], reverse=True)

        # 顶部半透明黑色背景条
        banner_h = 40 + min(len(sorted_alerts), 5) * 34
        overlay = frame.copy()
        cv2.rectangle(overlay, (0, 0), (w, banner_h), (0, 0, 0), -1)
        cv2.addWeighted(overlay, 0.6, frame, 0.4, 0, frame)

        # 标题
        top = 28
        title = "!! LIVE ALERT !!"
        (tw, th), _ = cv2.getTextSize(title, cv2.FONT_HERSHEY_SIMPLEX, 0.8, 2)
        cv2.putText(frame, title, (12, top), cv2.FONT_HERSHEY_SIMPLEX, 0.8, (255, 255, 255), 2)

        # 逐条绘制预警（最多显示5条）
        top += 4
        for i, (obj_id, cls_label, alert_type, level, timestamp) in enumerate(sorted_alerts[:5]):
            top += 30
            color_bgr, _ = level_color_map.get(level, ((255, 255, 255), ""))
            time_str = datetime.fromtimestamp(timestamp).strftime("%H:%M:%S")
            atype_en = type_en.get(alert_type, alert_type)
            text = f"[{level_text_map.get(level, '?')}] {atype_en}: {cls_label}-{obj_id}  {time_str}"

            # 闪烁效果：三级危险用脉冲
            if level == 3:
                pulse = int(np.abs(np.sin(time.time() * 4)) * 255)
                display_color = (pulse, pulse // 3, 255 - pulse // 3) if pulse > 100 else (0, 0, 255)
            else:
                display_color = color_bgr

            cv2.putText(frame, text, (12, top), cv2.FONT_HERSHEY_SIMPLEX, 0.55, display_color, 2)

        # 三级危险 - 屏幕边缘红色闪烁边框
        has_level3 = any(a[3] == 3 for a in new_alerts)
        if has_level3:
            flash = int(np.abs(np.sin(time.time() * 6)) * 180) + 60
            cv2.rectangle(frame, (4, 4), (w - 4, h - 4), (0, 0, flash), 5)
            flash_text = "!! DANGER !!"
            (ftw, fth), _ = cv2.getTextSize(flash_text, cv2.FONT_HERSHEY_SIMPLEX, 1.2, 3)
            fx = (w - ftw) // 2
            fy = h - 36
            cv2.rectangle(frame, (fx - 12, fy - fth - 12), (fx + ftw + 12, fy + 12), (0, 0, 0), -1)
            cv2.putText(frame, flash_text, (fx, fy), cv2.FONT_HERSHEY_SIMPLEX, 1.2, (0, 0, flash), 3)

    def apply_privacy_mask(self, frame):
        try:
            face_cascade = cv2.CascadeClassifier(cv2.data.haarcascades + 'haarcascade_frontalface_default.xml')
            gray = cv2.cvtColor(frame, cv2.COLOR_BGR2GRAY)
            faces = face_cascade.detectMultiScale(gray, 1.1, 4)
            for (x, y, w, h) in faces:
                roi = frame[y:y+h, x:x+w]
                roi = cv2.GaussianBlur(roi, (51, 51), 30)
                frame[y:y+h, x:x+w] = roi
        except:
            pass
        return frame

    def check_intrusion(self, x1, y1, x2, y2, obj_id=0, label=""):
        for zone in self.intrusion_zones:
            zx1, zy1, zx2, zy2 = zone
            if x1 < zx2 and x2 > zx1 and y1 < zy2 and y2 > zy1:
                self.intrusion_alerts.append(f"警告：{label}-{obj_id} 进入受限区域！")
                return True, (obj_id, label, "区域入侵", 2, time.time())
        return False, None

    # ---------- 预警UI更新 ----------
    def update_alert_ui(self, new_alerts):
        """更新预警列表，存储到历史并刷新（尊重手动过滤设置）"""
        level_text = {1: "注意", 2: "警告", 3: "危险"}

        for (obj_id, cls_label, alert_type, level, timestamp) in new_alerts:
            time_str = datetime.fromtimestamp(timestamp).strftime("%H:%M:%S")
            self.alert_history.append({
                "obj_id": obj_id, "label": cls_label,
                "type": alert_type, "level": level, "level_name": level_text[level],
                "time": time_str, "timestamp": timestamp
            })
            if len(self.alert_history) > 500:
                self.alert_history = self.alert_history[-300:]

        # 通过过滤刷新列表
        self._apply_alert_filter()

    # ---------- 显示与交互 ----------
    def update_stats_table(self):
        self.stats_table.setRowCount(len(self.class_stats))
        for row, (label, count) in enumerate(self.class_stats.items()):
            self.stats_table.setItem(row, 0, QTableWidgetItem(label))
            self.stats_table.setItem(row, 1, QTableWidgetItem(str(count)))

    def display_image(self, img_array):
        h, w, ch = img_array.shape
        bytes_per_line = ch * w
        # .copy() 确保 QImage 持有独立数据，不依赖 numpy 缓冲区
        qimg = QImage(img_array.data, w, h, bytes_per_line, QImage.Format_RGB888).copy()
        pixmap = QPixmap.fromImage(qimg)
        label_size = self.image_label.size()
        if label_size.width() < 10:
            label_size = pixmap.size()
        self.image_label.setPixmap(pixmap.scaled(label_size, Qt.KeepAspectRatio, Qt.SmoothTransformation))
        self.image_label.update()

    def set_qwen_api_key(self, api_key):
        self.qwen_analyzer.api_key = api_key
        QMessageBox.information(self, "API Key", "千问 API Key 已设置")

    def generate_report(self):
        # 如果没有设置 API Key，弹出输入框
        if not self.qwen_analyzer.api_key:
            key, ok = QInputDialog.getText(
                self, "设置千问 API Key",
                "请输入千问 (DashScope) API Key:\n(后续可在 生成报告 时重新设置)",
                QLineEdit.Password
            )
            if ok and key.strip():
                self.qwen_analyzer.api_key = key.strip()
            else:
                # 取消 或 空输入 → 直接生成不含AI分析的基础报告
                self._build_word_report("", "", "未配置千问 API，本报告仅含基础数据。")
                return

        # 收集数据
        total_targets = sum(self.class_stats.values())
        stats_summary = ", ".join(f"{k}:{v}" for k, v in sorted(self.class_stats.items(), key=lambda x: -x[1])[:15])
        alert_summary = ""
        alert_levels = {}
        if self.alert_history:
            for a in self.alert_history:
                alert_levels[a["level_name"]] = alert_levels.get(a["level_name"], 0) + 1
            alert_summary = ", ".join(f"{k}:{v}次" for k, v in alert_levels.items())
            type_count = defaultdict(int)
            for a in self.alert_history:
                type_count[a["type"]] += 1
            alert_summary += " | 类型: " + ", ".join(f"{k}:{v}" for k, v in type_count.items())

        # 构建千问分析 prompt
        qwen_prompt = f"""请根据以下目标检测系统的数据，生成一份专业的分析报告：

【检测统计】
- 检测目标总数: {total_targets}
- 检测类别分布: {stats_summary if stats_summary else "无"}
- 检测模式: {self.detection_mode}
- 置信度阈值: {self.confidence_threshold:.2f}

【预警统计】
- 预警总数: {len(self.alert_history)}
- 按等级分布: {alert_summary if alert_summary else "无预警"}
- 入侵区域数量: {len(self.intrusion_zones)}

请生成一份结构完整的报告，包含：
1. 总体概述（检测系统运行状态总结）
2. 检测数据分析（各类目标分布情况解读）
3. 安全预警评估（预警等级分布和风险判断）
4. 改进建议（基于数据的安全和管理建议）"""

        self.progress = QProgressDialog("正在调用千问大模型生成报告...", "取消", 0, 0, self)
        self.progress.setWindowTitle("生成报告")
        self.progress.setWindowModality(Qt.WindowModal)
        self.progress.show()

        # ---- 线程安全：通过 pyqtSignal 把结果投递回主线程 ----
        bridge = self.qwen_analyzer._bridge
        try:
            bridge.finished.disconnect()
        except TypeError:
            pass
        try:
            bridge.error.disconnect()
        except TypeError:
            pass

        def on_analysis(analysis_text):
            """工作线程回调 — 通过信号安全投递到主线程"""
            bridge.finished.emit(stats_summary, alert_summary, analysis_text)

        def on_report_ready(s_summary, a_summary, analysis_text):
            """主线程槽 — 安全执行 GUI 操作"""
            self.progress.close()
            self._build_word_report(s_summary, a_summary, analysis_text)

        def on_report_error(msg):
            self.progress.close()
            QMessageBox.warning(self, "报告生成失败", msg)

        bridge.finished.connect(on_report_ready)
        bridge.error.connect(on_report_error)

        self.qwen_analyzer.analyze(qwen_prompt, callback=on_analysis)

    def _build_word_report(self, stats_summary, alert_summary, analysis_text):
        """生成 Word 文档"""
        try:
            from docx import Document
            from docx.shared import Inches, Pt, Cm, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.table import WD_TABLE_ALIGNMENT
        except ImportError:
            QMessageBox.critical(self, "缺少依赖", "请安装 python-docx: pip install python-docx")
            return

        try:
            doc = Document()
            style = doc.styles['Normal']
            font = style.font
            font.name = 'Microsoft YaHei'
            font.size = Pt(11)

            # 标题
            title = doc.add_heading('智能视觉检测分析报告', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            doc.add_paragraph(f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            doc.add_paragraph(f"操作用户: {self.username}")
            doc.add_paragraph(f"检测模式: {self.detection_mode} | 置信度阈值: {self.confidence_threshold:.2f}")
            doc.add_paragraph("")

            # 千问分析内容
            doc.add_heading('一、AI 智能分析', 1)
            for para_text in analysis_text.split('\n'):
                if para_text.strip():
                    cleaned = para_text.strip()
                    if cleaned.startswith('###'):
                        doc.add_heading(cleaned.replace('###', '').strip(), 2)
                    elif cleaned.startswith('##'):
                        doc.add_heading(cleaned.replace('##', '').strip(), 1)
                    else:
                        doc.add_paragraph(cleaned)

            doc.add_page_break()

            # 检测数据汇总
            doc.add_heading('二、检测数据汇总', 1)
            doc.add_paragraph(f"检测目标总数: {sum(self.class_stats.values())}")

            table = doc.add_table(rows=1, cols=3)
            table.style = 'Light List Accent 1'
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            hdr = table.rows[0].cells
            hdr[0].text = '目标类别'
            hdr[1].text = '数量'
            hdr[2].text = '占比'
            for label, count in sorted(self.class_stats.items(), key=lambda x: -x[1]):
                row = table.add_row().cells
                row[0].text = label
                row[1].text = str(count)
                total = max(sum(self.class_stats.values()), 1)
                row[2].text = f"{count / total * 100:.1f}%"

            doc.add_paragraph("")

            # 预警数据
            doc.add_heading('三、安全预警数据', 1)
            doc.add_paragraph(f"预警总数: {len(self.alert_history)}")
            doc.add_paragraph(f"预警摘要: {alert_summary}")

            if self.alert_history:
                alert_table = doc.add_table(rows=1, cols=4)
                alert_table.style = 'Light List Accent 1'
                alert_table.alignment = WD_TABLE_ALIGNMENT.CENTER
                ahdr = alert_table.rows[0].cells
                ahdr[0].text = '等级'
                ahdr[1].text = '目标'
                ahdr[2].text = '类型'
                ahdr[3].text = '时间'
                for alert in self.alert_history[-50:]:
                    arow = alert_table.add_row().cells
                    arow[0].text = alert["level_name"]
                    arow[1].text = f'{alert["label"]}-{alert["obj_id"]}'
                    arow[2].text = alert["type"]
                    arow[3].text = alert["time"]

            # 保存
            save_path, _ = QFileDialog.getSaveFileName(
                self, "保存分析报告", f"检测报告_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                "Word 文档 (*.docx)"
            )
            if save_path:
                doc.save(save_path)
                QMessageBox.information(self, "报告生成成功", f"报告已保存至:\n{save_path}")

        except Exception as e:
            QMessageBox.critical(self, "报告生成失败", f"生成 Word 文档时出错:\n{str(e)}")

    # ---------- 图片/视频控制 ----------
    def release_resources(self):
        """停止视频/摄像头，释放资源"""
        with QMutexLocker(self.mutex):
            if self.timer.isActive():
                self.timer.stop()
                try:
                    self.timer.timeout.disconnect()
                except TypeError:
                    pass
            if self.cap:
                self.cap.release()
                self.cap = None
            self.status_label.setText("检测已停止")
            self.stop_button.setEnabled(False)

    def _stop_timer_if_running(self):
        """仅停止计时器，不清除画面（图片检测用）"""
        if self.timer.isActive():
            self.timer.stop()
            try:
                self.timer.timeout.disconnect()
            except TypeError:
                pass
        if self.cap:
            self.cap.release()
            self.cap = None

    def select_image(self):
        self._stop_timer_if_running()
        file_path, _ = QFileDialog.getOpenFileName(self, "选择图片", "", "Images (*.png *.jpg *.jpeg)")
        if file_path:
            img = cv2.imread(file_path)
            if img is None:
                QMessageBox.critical(self, "图片错误", "无法读取图片文件")
                return
            # 保持宽高比，限制最大边不超过 1920
            h, w = img.shape[:2]
            max_dim = 1920
            if max(h, w) > max_dim:
                scale = max_dim / max(h, w)
                img = cv2.resize(img, (int(w * scale), int(h * scale)))
            result_img = self.detect_objects(img, is_image=True)
            self.last_frame_rgb = cv2.cvtColor(result_img, cv2.COLOR_BGR2RGB)
            self.display_image(self.last_frame_rgb)
            self.status_label.setText(f"图片检测完成 · 总数: {sum(self.class_stats.values())}")
            default_path = os.path.splitext(file_path)[0] + '_detected' + os.path.splitext(file_path)[1]
            save_path, _ = QFileDialog.getSaveFileName(self, "保存检测结果", default_path, "PNG 图片 (*.png);;JPEG 图片 (*.jpg *.jpeg)")
            if save_path:
                cv2.imwrite(save_path, result_img)
                targets = self._get_detected_targets()
                self.save_detection_to_db(save_path, targets, sum(self.class_stats.values()))

    def select_video(self):
        self.release_resources()
        file_path, _ = QFileDialog.getOpenFileName(self, "选择视频", "", "Videos (*.mp4 *.avi *.mov)")
        if file_path:
            self.cap = cv2.VideoCapture(file_path)
            if not self.cap.isOpened():
                QMessageBox.critical(self, "视频错误", "无法打开视频文件")
                return
            self.run_video_detection()

    def detect_camera(self):
        self.release_resources()
        self.cap = cv2.VideoCapture(0)
        if not self.cap.isOpened():
            QMessageBox.critical(self, "摄像头错误", "无法打开摄像头")
            return
        self.run_video_detection()

    def run_video_detection(self):
        try:
            self.timer.timeout.disconnect()
        except:
            pass
        self.timer.timeout.connect(self.read_frame)
        self.timer.start(33)  # 30 FPS
        self.stop_button.setEnabled(True)

    def read_frame(self):
        ret, frame = self.cap.read()
        if not ret:
            self.timer.stop()
            self.cap.release()
            self.cap = None
            self.status_label.setText("视频流已结束")
            return
        result_frame = self.detect_objects(frame)
        self.last_frame_rgb = cv2.cvtColor(result_frame, cv2.COLOR_BGR2RGB)
        self.display_image(self.last_frame_rgb)
        self.status_label.setText(f"检测中 · 总数: {sum(self.class_stats.values())}")

    def stop_detection(self):
        self.release_resources()

    def save_detection_to_db(self, image_path, target, count):
        """保存检测记录到数据库，target 为实际检测到的类别列表"""
        try:
            conn = db_connect()
            cursor = conn.cursor()
            cursor.execute(
                "INSERT INTO detections (username, image_path, target_type, count, file_name) VALUES (%s, %s, %s, %s, %s)",
                (self.username, image_path, target, count, os.path.basename(image_path))
            )
            conn.commit()
            conn.close()
        except Exception as e:
            print("保存失败：", e)

    def _get_detected_targets(self):
        """获取当前检测到的目标类型汇总，如 'person(3), car(2), truck(1)'"""
        if not self.class_stats:
            return "无"
        items = sorted(self.class_stats.items(), key=lambda x: -x[1])[:10]
        return ", ".join(f"{label}({cnt})" for label, cnt in items)

    def show_history(self):
        history = HistoryWindow(self.username, self)
        history.exec_()
